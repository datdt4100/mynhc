#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script tạo file Word hướng dẫn sử dụng mynhc (ClassReg)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─────────────────────────────────────────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────────────────────────────────────────

def set_page_margins(doc, top=2.5, bottom=2.5, left=2.5, right=2.5):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)

def set_font(run, name="Times New Roman", size=13, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Ensure East-Asian font also set
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)

def add_heading(doc, text, level=1, center=False):
    style_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3'}
    p = doc.add_paragraph(style=style_map.get(level, 'Heading 1'))
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    else:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x1F, 0x74, 0x89)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Ensure East-Asian font
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), "Times New Roman")
    rFonts.set(qn('w:hAnsi'), "Times New Roman")
    rFonts.set(qn('w:cs'), "Times New Roman")
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)
    return p

def add_body(doc, text, bold=False, italic=False, indent=False, center=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic, color=color)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_numbered_step(doc, number, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    set_font(run)

def add_bullet(doc, text, indent_level=1):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(indent_level)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run)

def shade_paragraph(p, hex_color="FFF2CC"):
    """Apply background shading to a paragraph."""
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def add_note_box(doc, text, title="Lưu ý", color="FFF2CC", title_color=(0xBF, 0x8F, 0x00)):
    """Add a styled note box."""
    p_title = doc.add_paragraph()
    p_title.paragraph_format.left_indent  = Cm(0.5)
    p_title.paragraph_format.right_indent = Cm(0.5)
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after  = Pt(0)
    run_title = p_title.add_run(f"  {title}")
    set_font(run_title, bold=True, size=12, color=title_color)
    shade_paragraph(p_title, color)

    p_body = doc.add_paragraph()
    p_body.paragraph_format.left_indent  = Cm(0.5)
    p_body.paragraph_format.right_indent = Cm(0.5)
    p_body.paragraph_format.space_before = Pt(0)
    p_body.paragraph_format.space_after  = Pt(6)
    run_body = p_body.add_run(f"  {text}")
    set_font(run_body, size=12, italic=True)
    shade_paragraph(p_body, color)
    return p_body

def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            set_font(run, bold=True, size=12)
        # Blue background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2E74B5')
        tcPr.append(shd)
        # White font
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(cell_text)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in cell.paragraphs[0].runs:
                set_font(run, size=12)
            # Alternate row shading
            if r_idx % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'DEEAF1')
                tcPr.append(shd)

    # Set column widths
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[i])

    doc.add_paragraph()  # spacing after table
    return table

def add_header_footer(doc):
    """Add header and footer to all sections."""
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run("Hướng dẫn sử dụng mynhc — Trường THPT Nguyễn Hữu Cầu")
        set_font(run, size=11, italic=True, color=(0x44, 0x72, 0xC4))
        # Header bottom border
        pPr = hp._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '4472C4')
        pBdr.append(bottom)
        pPr.append(pBdr)

        # Footer with page number
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_prefix = fp.add_run("Trang ")
        set_font(run_prefix, size=11)
        # Page number field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = ' PAGE '
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run_pg = fp.add_run()
        run_pg._r.append(fldChar1)
        run_pg._r.append(instrText)
        run_pg._r.append(fldChar2)
        run_pg._r.append(fldChar3)
        set_font(run_pg, size=11)

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)

import docx.enum.text

# ─────────────────────────────────────────────────────────────────────────────
# PAGE SETUP
# ─────────────────────────────────────────────────────────────────────────────
set_page_margins(doc, 2.5, 2.5, 3.0, 2.0)

# ─────────────────────────────────────────────────────────────────────────────
# TRANG BÌA
# ─────────────────────────────────────────────────────────────────────────────
# Top info
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
run = p.add_run("SỞ GIÁO DỤC VÀ ĐÀO TẠO TỈNH TIỀN GIANG")
set_font(run, bold=True, size=13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("TRƯỜNG THPT NGUYỄN HỮU CẦU")
set_font(run, bold=True, size=14, color=(0x1F, 0x49, 0x7D))

# Divider
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("─" * 40)
set_font(run, size=12)

# Spacing
for _ in range(4):
    doc.add_paragraph()

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("HỆ THỐNG ĐĂNG KÝ LỚP HỌC")
set_font(run, bold=True, size=22, color=(0x1F, 0x49, 0x7D))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("mynhc")
set_font(run, bold=True, size=30, color=(0x2E, 0x74, 0xB5))

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("HƯỚNG DẪN SỬ DỤNG")
set_font(run, bold=True, size=18, color=(0xBF, 0x8F, 0x00))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Dành cho Quản trị viên – Giáo viên – Học sinh")
set_font(run, italic=True, size=13, color=(0x40, 0x40, 0x40))

for _ in range(6):
    doc.add_paragraph()

# Year
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Năm học 2026 – 2027")
set_font(run, bold=True, size=14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Phiên bản 1.1  |  Cập nhật: Tháng 8 năm 2026")
set_font(run, italic=True, size=11, color=(0x60, 0x60, 0x60))

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# MỤC LỤC
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "MỤC LỤC", 1, center=True)

toc_items = [
    ("1.", "TỔNG QUAN HỆ THỐNG", "3"),
    ("  1.1.", "Giới thiệu", "3"),
    ("  1.2.", "Các loại người dùng", "3"),
    ("  1.3.", "Quy trình hoạt động tổng quan", "4"),
    ("  1.4.", "Yêu cầu hệ thống", "4"),
    ("  1.5.", "Truy cập hệ thống", "4"),
    ("2.", "HƯỚNG DẪN ĐĂNG NHẬP", "5"),
    ("  2.1.", "Trang đăng nhập", "5"),
    ("  2.2.", "Đăng nhập cho Quản trị viên", "5"),
    ("  2.3.", "Đăng nhập cho Giáo viên", "5"),
    ("  2.4.", "Đăng nhập cho Học sinh", "6"),
    ("  2.5.", "Đăng xuất", "6"),
    ("3.", "HƯỚNG DẪN SỬ DỤNG – QUẢN TRỊ VIÊN", "7"),
    ("  3.1.", "Trang chủ quản trị", "7"),
    ("  3.2.", "Quản lý Giáo viên", "8"),
    ("  3.3.", "Quản lý Học sinh", "10"),
    ("  3.4.", "Quản lý Đăng ký Mở lớp (Giai đoạn 1)", "13"),
    ("  3.5.", "Quản lý Đăng ký Môn học (Giai đoạn 2)", "16"),
    ("4.", "HƯỚNG DẪN SỬ DỤNG – GIÁO VIÊN", "17"),
    ("  4.1.", "Tổng quan giao diện", "17"),
    ("  4.2.", "Xem thời khoá biểu", "17"),
    ("  4.3.", "Đăng ký mở lớp", "18"),
    ("  4.4.", "Xoá lớp đã đăng ký", "19"),
    ("5.", "HƯỚNG DẪN SỬ DỤNG – HỌC SINH", "20"),
    ("  5.1.", "Đặt mật khẩu lần đầu", "20"),
    ("  5.2.", "Tab Thời khoá biểu", "20"),
    ("  5.3.", "Tab Đăng ký môn học", "21"),
    ("  5.4.", "Đăng ký môn học", "22"),
    ("  5.5.", "Huỷ đăng ký môn học", "23"),
    ("  5.6.", "Xử lý trùng lịch", "23"),
    ("6.", "CÁC TÌNH HUỐNG THƯỜNG GẶP VÀ CÁCH XỬ LÝ", "24"),
    ("7.", "PHỤ LỤC", "26"),
]

toc_table = doc.add_table(rows=len(toc_items), cols=3)
toc_table.style = 'Table Grid'
col_widths_toc = [1.5, 12.5, 1.5]

for i, (num, title, page) in enumerate(toc_items):
    row = toc_table.rows[i]
    cells = row.cells
    cells[0].text = num
    cells[1].text = title
    cells[2].text = page
    for c_idx, cell in enumerate(cells):
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx < 2 else WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            bold = num.strip() and not num.startswith("  ")
            set_font(run, bold=bold and c_idx == 1, size=12)
        # Remove border for TOC
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border_el = OxmlElement(f'w:{side}')
            border_el.set(qn('w:val'), 'none')
            tcBorders.append(border_el)
        existing_border = tcPr.find(qn('w:tcBorders'))
        if existing_border is not None:
            tcPr.remove(existing_border)
        tcPr.append(tcBorders)

for row in toc_table.rows:
    for i, cell in enumerate(row.cells):
        cell.width = Cm(col_widths_toc[i])

doc.add_paragraph()
add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# PHẦN 1: TỔNG QUAN
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "1. TỔNG QUAN HỆ THỐNG", 1)

add_heading(doc, "1.1. Giới thiệu", 2)
add_body(doc, "Hệ thống mynhc (viết tắt từ Nguyễn Hữu Cầu) là nền tảng đăng ký lớp học trực tuyến được xây dựng riêng cho Trường THPT Nguyễn Hữu Cầu. Hệ thống giúp số hóa toàn bộ quy trình đăng ký môn học tự chọn, từ việc giáo viên đăng ký mở lớp đến học sinh chọn môn học theo thời khoá biểu.")
add_body(doc, "Hệ thống được phát triển trên nền tảng Flask (Python) kết hợp SQLite và giao diện Bootstrap 5, hỗ trợ truy cập trên máy tính và thiết bị di động.")

add_heading(doc, "1.2. Các loại người dùng", 2)
add_body(doc, "Hệ thống mynhc phục vụ ba nhóm người dùng với quyền hạn khác nhau:")

add_table(doc,
    ["Loại người dùng", "Mô tả", "Trang truy cập"],
    [
        ["Quản trị viên (Admin)", "Quản lý toàn bộ hệ thống: giáo viên, học sinh, điều phối các giai đoạn đăng ký, phân phòng học.", "/admin"],
        ["Giáo viên", "Đăng ký mở lớp theo lịch tuần, xem số học sinh đăng ký và phòng học được phân công.", "/teacher/schedule"],
        ["Học sinh", "Xem thời khoá biểu cá nhân, đăng ký và huỷ đăng ký các môn học.", "/student/schedule"],
    ],
    col_widths=[4.5, 9.0, 3.0]
)

add_heading(doc, "1.3. Quy trình hoạt động tổng quan", 2)
add_body(doc, "Hệ thống vận hành theo hai giai đoạn liên tiếp:")

add_note_box(doc,
    "Giai đoạn 1 → Giai đoạn 2: Quản trị viên phải hoàn tất phân phòng và nhập sĩ số tối đa "
    "trước khi mở Giai đoạn 2 cho học sinh đăng ký.",
    title="Quan trọng",
    color="FCE4D6",
    title_color=(0xC0, 0x50, 0x0)
)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run("Giai đoạn 1 – Giáo viên đăng ký mở lớp:")
set_font(run, bold=True, size=13)
add_bullet(doc, "Quản trị viên mở đăng ký Giai đoạn 1.")
add_bullet(doc, "Giáo viên đăng nhập và đăng ký các ô lịch (thứ, buổi, tiết, khối) muốn mở lớp.")
add_bullet(doc, "Quản trị viên xem danh sách lớp đã đăng ký, điền phòng học và sĩ số tối đa.")
add_bullet(doc, "Quản trị viên đóng Giai đoạn 1.")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run("Giai đoạn 2 – Học sinh đăng ký môn học:")
set_font(run, bold=True, size=13)
add_bullet(doc, "Quản trị viên mở đăng ký Giai đoạn 2.")
add_bullet(doc, "Học sinh đăng nhập và đăng ký các môn học theo khối của mình.")
add_bullet(doc, "Hệ thống kiểm tra trùng lịch tự động.")
add_bullet(doc, "Quản trị viên đóng đăng ký và xuất kết quả Excel.")

add_heading(doc, "1.4. Yêu cầu hệ thống", 2)
add_table(doc,
    ["Thành phần", "Yêu cầu"],
    [
        ["Trình duyệt web", "Google Chrome 90+, Firefox 88+, Microsoft Edge 90+, Safari 14+"],
        ["Kết nối mạng", "Kết nối nội bộ LAN trường hoặc Internet (tuỳ cách triển khai)"],
        ["Màn hình", "Tối thiểu độ phân giải 1024×768. Hỗ trợ màn hình di động"],
        ["JavaScript", "Phải bật JavaScript trên trình duyệt"],
    ],
    col_widths=[5.0, 11.5]
)

add_heading(doc, "1.5. Truy cập hệ thống", 2)
add_body(doc, "Hệ thống có thể chạy ở hai chế độ:")
add_bullet(doc, "Nội bộ (Local): http://localhost:5000  hoặc  http://<IP-máy-chủ>:5000")
add_bullet(doc, "Triển khai (Deploy): Theo địa chỉ do nhà trường công bố.")
add_note_box(doc,
    "Liên hệ bộ phận kỹ thuật của trường để biết địa chỉ URL chính thức khi hệ thống đã được triển khai.",
    title="Lưu ý",
    color="FFF2CC",
    title_color=(0xBF, 0x8F, 0x00)
)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# PHẦN 2: ĐĂNG NHẬP
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "2. HƯỚNG DẪN ĐĂNG NHẬP", 1)

add_heading(doc, "2.1. Trang đăng nhập", 2)
add_body(doc, "Tất cả người dùng đều đăng nhập qua cùng một trang tại địa chỉ /login. Giao diện trang đăng nhập gồm:")
add_table(doc,
    ["Trường / Nút", "Mô tả"],
    [
        ["Chọn loại người dùng", "Dropdown chọn một trong ba vai trò: Quản trị viên / Giáo viên / Học sinh"],
        ["Mã đăng nhập / MSSV / Tên đăng nhập", "Ô nhập mã định danh (hiển thị theo vai trò đã chọn)"],
        ["Mật khẩu", "Ô nhập mật khẩu (ẩn ký tự)"],
        ["Nút Đăng nhập", "Xác thực và chuyển hướng đến trang tương ứng"],
    ],
    col_widths=[6.0, 10.5]
)
add_body(doc, "[Hình ảnh: Giao diện trang đăng nhập /login với dropdown chọn vai trò]", italic=True, color=(0x70, 0x70, 0x70))

add_heading(doc, "2.2. Đăng nhập cho Quản trị viên", 2)
steps_admin_login = [
    "Truy cập URL hệ thống, ví dụ: http://localhost:5000/login",
    "Tại dropdown \"Loại người dùng\", chọn Quản trị viên.",
    "Nhập tên đăng nhập và mật khẩu quản trị viên.",
    "Nhấn nút Đăng nhập.",
    "Hệ thống chuyển hướng đến trang quản trị /admin.",
]
for i, s in enumerate(steps_admin_login, 1):
    add_numbered_step(doc, i, s)
add_note_box(doc,
    "Tài khoản quản trị viên được thiết lập sẵn khi cài đặt hệ thống. Liên hệ bộ phận kỹ thuật nếu quên mật khẩu.",
    title="Lưu ý",
)

add_heading(doc, "2.3. Đăng nhập cho Giáo viên", 2)
steps_gv_login = [
    "Truy cập URL hệ thống, ví dụ: http://localhost:5000/login",
    "Tại dropdown \"Loại người dùng\", chọn Giáo viên.",
    "Nhập Mã giáo viên (do quản trị viên cung cấp khi tạo tài khoản).",
    "Nhập mật khẩu (mặc định = mã giáo viên, trừ khi đã đổi).",
    "Nhấn nút Đăng nhập.",
    "Hệ thống chuyển hướng đến trang thời khoá biểu /teacher/schedule.",
]
for i, s in enumerate(steps_gv_login, 1):
    add_numbered_step(doc, i, s)

add_heading(doc, "2.4. Đăng nhập cho Học sinh", 2)
steps_hs_login = [
    "Truy cập URL hệ thống.",
    "Tại dropdown \"Loại người dùng\", chọn Học sinh.",
    "Nhập MSSV (Mã số sinh viên/học sinh).",
    "Nhập mật khẩu. Nếu là lần đầu tiên, nhập mật khẩu mặc định (MSSV) hoặc mật khẩu do quản trị viên reset.",
    "Nhấn nút Đăng nhập.",
    "Nếu đây là lần đăng nhập đầu tiên, hệ thống yêu cầu đặt mật khẩu mới (xem mục 5.1).",
    "Hệ thống chuyển hướng đến trang học sinh /student/schedule.",
]
for i, s in enumerate(steps_hs_login, 1):
    add_numbered_step(doc, i, s)

add_heading(doc, "2.5. Đăng xuất", 2)
add_body(doc, "Để đăng xuất khỏi hệ thống:")
add_numbered_step(doc, 1, "Nhấn vào nút Đăng xuất ở góc phải trên cùng của thanh điều hướng (header).")
add_numbered_step(doc, 2, "Hệ thống xoá phiên làm việc và chuyển về trang đăng nhập.")
add_note_box(doc,
    "Nên đăng xuất sau khi sử dụng, đặc biệt khi dùng máy tính chung để tránh người khác truy cập tài khoản.",
    title="Khuyến nghị bảo mật",
    color="E2EFDA",
    title_color=(0x37, 0x86, 0x10)
)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# PHẦN 3: QUẢN TRỊ VIÊN
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "3. HƯỚNG DẪN SỬ DỤNG – QUẢN TRỊ VIÊN", 1)
add_body(doc, "Quản trị viên là người điều phối toàn bộ hoạt động của hệ thống, từ quản lý tài khoản đến điều phối các giai đoạn đăng ký. Trang quản trị được truy cập tại /admin sau khi đăng nhập thành công.")

add_heading(doc, "3.1. Trang chủ quản trị", 2)
add_body(doc, "Ngay sau khi đăng nhập, quản trị viên được chuyển đến trang chủ /admin. Trang này hiển thị:")
add_table(doc,
    ["Thành phần", "Nội dung hiển thị"],
    [
        ["Thẻ thống kê – Số giáo viên", "Tổng số tài khoản giáo viên trong hệ thống"],
        ["Thẻ thống kê – Số học sinh", "Tổng số tài khoản học sinh trong hệ thống"],
        ["Thẻ thống kê – Số lớp đã mở", "Tổng số lớp giáo viên đã đăng ký mở"],
        ["Thẻ thống kê – Lượt đăng ký", "Tổng số lượt học sinh đã đăng ký môn học"],
        ["Trạng thái Giai đoạn 1", "Đang mở / Đã đóng (với nút Mở/Đóng)"],
        ["Trạng thái Giai đoạn 2", "Đang mở / Đã đóng (với nút Mở/Đóng)"],
        ["Menu điều hướng", "Các liên kết đến: Quản lý GV, Quản lý HS, Đăng ký mở lớp, Đăng ký môn học"],
    ],
    col_widths=[6.5, 10.0]
)
add_body(doc, "[Hình ảnh: Trang chủ quản trị với 4 thẻ thống kê và bảng điều khiển giai đoạn]", italic=True, color=(0x70, 0x70, 0x70))

p = doc.add_paragraph()
run = p.add_run("Các bước thao tác điều phối giai đoạn từ trang chủ:")
set_font(run, bold=True)

add_numbered_step(doc, 1, "Đăng nhập với tài khoản quản trị viên.")
add_numbered_step(doc, 2, "Xem tổng quan thống kê trên 4 thẻ số liệu.")
add_numbered_step(doc, 3, "Khi muốn bắt đầu giai đoạn đăng ký giáo viên: nhấn nút Mở tại mục Giai đoạn 1.")
add_numbered_step(doc, 4, "Khi muốn kết thúc đăng ký giáo viên: nhấn nút Đóng tại mục Giai đoạn 1.")
add_numbered_step(doc, 5, "Khi muốn bắt đầu đăng ký cho học sinh: nhấn nút Mở tại mục Giai đoạn 2.")
add_numbered_step(doc, 6, "Sử dụng menu điều hướng để đến các trang quản lý chi tiết.")

add_heading(doc, "3.2. Quản lý Giáo viên", 2)
add_body(doc, "Trang Quản lý Giáo viên cho phép thêm, sửa, xoá thông tin tài khoản giáo viên. Truy cập từ menu Admin → Quản lý Giáo viên.")

add_heading(doc, "3.2.1. Xem danh sách giáo viên", 3)
add_body(doc, "Trang hiển thị bảng danh sách tất cả giáo viên gồm các cột:")
add_table(doc,
    ["Cột", "Mô tả"],
    [
        ["STT", "Số thứ tự"],
        ["Mã giáo viên", "Mã định danh dùng để đăng nhập (CCCD)"],
        ["Họ và tên", "Tên đầy đủ của giáo viên"],
        ["Môn dạy", "Môn học giáo viên phụ trách"],
        ["Email", "Địa chỉ email (nếu có)"],
        ["Trạng thái TK", "Đã kích hoạt / Chưa kích hoạt (màu xanh / xám)"],
        ["Thao tác", "Nút icon Sửa, Xoá, Khôi phục MK (chỉ hiện khi tài khoản đã kích hoạt)"],
    ],
    col_widths=[1.2, 3.0, 3.5, 2.8, 3.5, 3.0, 1.5]
)
add_body(doc, "Ô tìm kiếm phía trên bảng cho phép lọc giáo viên theo Họ tên hoặc Mã GV (CCCD) ngay lập tức mà không cần tải lại trang.")
add_note_box(doc,
    "Tài khoản chưa kích hoạt (is_first_login = 1) không hiển thị nút Khôi phục MK và Reset. "
    "Giáo viên phải đăng nhập lần đầu và đặt mật khẩu mới thì tài khoản mới được coi là đã kích hoạt.",
    title="Lưu ý",
)

add_heading(doc, "3.2.2. Thêm giáo viên mới", 3)
add_numbered_step(doc, 1, "Nhấn nút Thêm giáo viên (màu xanh lá) phía trên bảng danh sách.")
add_numbered_step(doc, 2, "Điền đầy đủ thông tin trong form: Mã giáo viên, Họ và tên, Môn dạy, Email (tuỳ chọn).")
add_numbered_step(doc, 3, "Nhập mật khẩu ban đầu cho giáo viên (mặc định có thể đặt trùng mã GV).")
add_numbered_step(doc, 4, "Nhấn Lưu để tạo tài khoản.")
add_numbered_step(doc, 5, "Thông báo xanh lá xác nhận tạo thành công sẽ xuất hiện ở đầu trang.")

add_heading(doc, "3.2.3. Sửa thông tin giáo viên", 3)
add_numbered_step(doc, 1, "Tìm giáo viên cần sửa trong bảng danh sách.")
add_numbered_step(doc, 2, "Nhấn nút Sửa (biểu tượng bút chì) ở cột Thao tác.")
add_numbered_step(doc, 3, "Cập nhật thông tin trong form hiện ra.")
add_numbered_step(doc, 4, "Nhấn Lưu thay đổi để xác nhận.")

add_heading(doc, "3.2.4. Xoá giáo viên", 3)
add_numbered_step(doc, 1, "Nhấn nút Xoá (biểu tượng thùng rác, màu đỏ) ở cột Thao tác của giáo viên cần xoá.")
add_numbered_step(doc, 2, "Hộp thoại xác nhận xuất hiện. Nhấn Xác nhận để tiến hành xoá.")
add_note_box(doc,
    "Xoá giáo viên sẽ xoá toàn bộ lớp học mà giáo viên đó đã đăng ký mở. "
    "Hãy chắc chắn trước khi thực hiện thao tác này.",
    title="Cảnh báo",
    color="FCE4D6",
    title_color=(0xC0, 0x50, 0x0)
)

add_heading(doc, "3.2.5. Khôi phục mật khẩu giáo viên", 3)
add_body(doc, "Khi giáo viên quên mật khẩu, quản trị viên tạo mật khẩu tạm thời và buộc đổi khi đăng nhập lần sau:")
add_numbered_step(doc, 1, "Tìm giáo viên đã kích hoạt trong danh sách.")
add_numbered_step(doc, 2, "Nhấn nút Khôi phục MK (biểu tượng chìa khoá, màu cam) ở cột Thao tác.")
add_numbered_step(doc, 3, "Hộp thoại xác nhận xuất hiện. Nhấn Xác nhận.")
add_numbered_step(doc, 4, "Hệ thống tạo mật khẩu tạm thời ngẫu nhiên 12 ký tự và hiển thị trong modal màu vàng.")
add_numbered_step(doc, 5, "Sao chép mật khẩu tạm thời (có nút Copy) và cung cấp cho giáo viên.")
add_numbered_step(doc, 6, "Khi giáo viên đăng nhập bằng mật khẩu tạm thời, hệ thống tự động chuyển đến trang Đổi mật khẩu bắt buộc.")
add_note_box(doc,
    "Mật khẩu tạm thời CHỈ hiển thị một lần duy nhất trong modal ngay sau khi tạo. "
    "Hệ thống không lưu lại mật khẩu này ở dạng văn bản thuần, nên hãy sao chép ngay trước khi đóng modal.",
    title="Quan trọng",
    color="FCE4D6",
    title_color=(0xC0, 0x50, 0x0)
)

add_heading(doc, "3.2.6. Nhập danh sách giáo viên từ Excel", 3)
add_numbered_step(doc, 1, "Nhấn nút Nhập Excel (biểu tượng file) ở phía trên bảng.")
add_numbered_step(doc, 2, "Tải về file mẫu Excel để biết định dạng cột cần nhập.")
add_numbered_step(doc, 3, "Điền dữ liệu vào file mẫu theo đúng cột: Mã GV, Họ tên, Môn dạy, Email, Mật khẩu.")
add_numbered_step(doc, 4, "Upload file Excel đã điền lên hệ thống.")
add_numbered_step(doc, 5, "Hệ thống xử lý và báo cáo số bản ghi thành công / lỗi.")
add_body(doc, "[Hình ảnh: Modal nhập Excel với bảng kết quả import]", italic=True, color=(0x70, 0x70, 0x70))

add_heading(doc, "3.3. Quản lý Học sinh", 2)
add_body(doc, "Trang Quản lý Học sinh cung cấp đầy đủ công cụ để quản lý tài khoản học sinh, bao gồm thêm/sửa/xoá và reset mật khẩu. Truy cập từ menu Admin → Quản lý Học sinh.")

add_heading(doc, "3.3.1. Xem danh sách học sinh", 3)
add_table(doc,
    ["Cột", "Mô tả"],
    [
        ["STT", "Số thứ tự"],
        ["CCCD", "Số CCCD, dùng để đăng nhập"],
        ["Họ và tên", "Tên đầy đủ học sinh"],
        ["Khối", "Khối lớp học sinh đang học (10, 11 hoặc 12)"],
        ["Lớp", "Tên lớp học (ví dụ: 10A1, 11B2...)"],
        ["Email", "Địa chỉ email học sinh (nếu có)"],
        ["Trạng thái TK", "Đã kích hoạt / Chưa kích hoạt (màu xanh / xám)"],
        ["Thao tác", "Nút icon Sửa, Xoá, Khôi phục MK (chỉ hiện khi tài khoản đã kích hoạt)"],
    ],
    col_widths=[1.0, 3.0, 3.5, 1.5, 2.0, 3.5, 2.5, 1.5]
)
add_body(doc, "Ô tìm kiếm phía trên bảng cho phép lọc học sinh theo Họ tên hoặc CCCD ngay lập tức.")

add_heading(doc, "3.3.2. Thêm học sinh mới", 3)
add_numbered_step(doc, 1, "Nhấn nút Thêm học sinh phía trên bảng danh sách.")
add_numbered_step(doc, 2, "Điền: MSSV, Họ tên, Khối (10/11/12), Lớp.")
add_numbered_step(doc, 3, "Mật khẩu mặc định sẽ là MSSV; học sinh cần đổi mật khẩu trong lần đầu đăng nhập.")
add_numbered_step(doc, 4, "Nhấn Lưu để tạo tài khoản.")

add_heading(doc, "3.3.3. Khôi phục mật khẩu học sinh", 3)
add_body(doc, "Khi học sinh quên mật khẩu, quản trị viên tạo mật khẩu tạm thời ngẫu nhiên:")
add_numbered_step(doc, 1, "Tìm học sinh đã kích hoạt trong danh sách.")
add_numbered_step(doc, 2, "Nhấn nút Khôi phục MK (biểu tượng chìa khoá, màu cam) ở cột Thao tác.")
add_numbered_step(doc, 3, "Hộp thoại xác nhận xuất hiện. Nhấn Xác nhận.")
add_numbered_step(doc, 4, "Hệ thống sinh mật khẩu tạm thời ngẫu nhiên 12 ký tự và hiển thị trong modal màu vàng.")
add_numbered_step(doc, 5, "Sao chép mật khẩu tạm thời (có nút Copy) và cung cấp cho học sinh.")
add_numbered_step(doc, 6, "Học sinh đăng nhập bằng mật khẩu tạm thời → hệ thống tự chuyển đến trang Đổi mật khẩu bắt buộc.")
add_note_box(doc,
    "Mật khẩu tạm thời chỉ hiển thị một lần. Hãy sao chép ngay trước khi đóng modal.",
    title="Quan trọng",
    color="FCE4D6",
    title_color=(0xC0, 0x50, 0x0)
)

add_heading(doc, "3.3.4. Nhập danh sách học sinh từ Excel", 3)
add_body(doc, "Thao tác tương tự nhập danh sách giáo viên. File Excel mẫu gồm các cột: MSSV, Họ tên, Khối, Lớp.")
add_numbered_step(doc, 1, "Nhấn nút Nhập Excel.")
add_numbered_step(doc, 2, "Tải file mẫu, điền dữ liệu học sinh.")
add_numbered_step(doc, 3, "Upload file và chờ hệ thống xử lý.")
add_numbered_step(doc, 4, "Xem báo cáo kết quả nhập: số bản ghi thành công, số lỗi và lý do lỗi.")
add_body(doc, "[Hình ảnh: Trang quản lý học sinh với bảng danh sách và bộ lọc theo khối/lớp]", italic=True, color=(0x70, 0x70, 0x70))

add_heading(doc, "3.4. Quản lý Đăng ký Mở lớp (Giai đoạn 1)", 2)
add_body(doc, "Đây là bước trung tâm của Giai đoạn 1. Quản trị viên vào trang /admin/class-reg để:")
add_bullet(doc, "Mở/đóng đăng ký cho giáo viên.")
add_bullet(doc, "Xem danh sách lớp theo từng giáo viên.")
add_bullet(doc, "Điền phòng học và sĩ số tối đa cho từng lớp.")
add_bullet(doc, "Xuất/nhập file Excel phân phòng.")

add_heading(doc, "3.4.1. Mở và đóng Giai đoạn 1", 3)
add_numbered_step(doc, 1, "Truy cập Admin → Quản lý đăng ký mở lớp hoặc URL /admin/class-reg.")
add_numbered_step(doc, 2, "Nhấn nút Mở đăng ký GV để giáo viên có thể đăng ký lịch mở lớp.")
add_numbered_step(doc, 3, "Khi đã thu thập đủ đăng ký, nhấn Đóng đăng ký GV để kết thúc Giai đoạn 1.")
add_note_box(doc,
    "Sau khi đóng Giai đoạn 1, giáo viên không thể thêm hoặc xoá lớp đã đăng ký. "
    "Hãy thông báo cho giáo viên trước khi đóng.",
    title="Lưu ý",
)

add_heading(doc, "3.4.2. Xem danh sách lớp theo giáo viên", 3)
add_body(doc, "Trang hiển thị danh sách các giáo viên. Mỗi giáo viên có thể mở rộng (accordion) để xem các lớp đã đăng ký:")
add_table(doc,
    ["Thông tin", "Mô tả"],
    [
        ["Môn học", "Tên môn giáo viên sẽ dạy"],
        ["Khối", "Khối học sinh dự kiến (10/11/12)"],
        ["Thứ / Buổi / Tiết", "Lịch dạy cụ thể trong tuần"],
        ["Phòng học", "Ô nhập phòng (để trống nếu chưa phân công)"],
        ["Sĩ số tối đa", "Số học sinh tối đa được đăng ký vào lớp"],
        ["Trạng thái", "Cảnh báo đỏ nếu thiếu phòng hoặc sĩ số"],
    ],
    col_widths=[3.0, 8.5, 5.0]
)
add_body(doc, "[Hình ảnh: Accordion danh sách lớp theo giáo viên, ô phòng học và sĩ số tối đa]", italic=True, color=(0x70, 0x70, 0x70))

add_heading(doc, "3.4.3. Điền phòng học và sĩ số tối đa", 3)
add_body(doc, "Với mỗi lớp học trong danh sách:")
add_numbered_step(doc, 1, "Click vào tên giáo viên để mở rộng danh sách lớp của GV đó.")
add_numbered_step(doc, 2, "Tại cột Phòng học, click vào ô trống và nhập số phòng (ví dụ: A201, Hội trường...).")
add_numbered_step(doc, 3, "Nhấn Enter hoặc click ra ngoài ô (blur) để lưu tự động — không cần nhấn nút Lưu riêng.")
add_numbered_step(doc, 4, "Tại cột Sĩ số tối đa, nhập số học sinh tối đa cho lớp đó.")
add_numbered_step(doc, 5, "Nhấn Enter hoặc click ra ngoài để lưu.")
add_numbered_step(doc, 6, "Lớp có viền đỏ = thiếu phòng hoặc thiếu sĩ số → cần bổ sung trước khi mở Giai đoạn 2.")
add_note_box(doc,
    "Tính năng auto-save (lưu tự động): dữ liệu được ghi ngay khi bạn rời khỏi ô nhập (sự kiện blur) "
    "hoặc nhấn Enter. Không cần nhấn nút Lưu tổng. Ô chuyển màu xanh nhạt để xác nhận đã lưu.",
    title="Tính năng Auto-save",
    color="E2EFDA",
    title_color=(0x37, 0x86, 0x10)
)

add_heading(doc, "3.4.4. Xuất Excel phân phòng", 3)
add_numbered_step(doc, 1, "Nhấn nút Xuất Excel phân phòng ở đầu trang.")
add_numbered_step(doc, 2, "File .xlsx được tải xuống máy, chứa toàn bộ thông tin: GV, môn, lịch, phòng, sĩ số.")
add_numbered_step(doc, 3, "Dùng file này để in lịch phân phòng hoặc gửi cho Ban Giám hiệu.")

add_heading(doc, "3.4.5. Upload file phân phòng", 3)
add_body(doc, "Nếu phòng học được phân công bằng file Excel bên ngoài, có thể upload ngược lại vào hệ thống:")
add_numbered_step(doc, 1, "Xuất file mẫu (nếu chưa có) theo định dạng chuẩn.")
add_numbered_step(doc, 2, "Điền thông tin phòng học và sĩ số vào file.")
add_numbered_step(doc, 3, "Nhấn Tải lên file phân phòng và chọn file đã chuẩn bị.")
add_numbered_step(doc, 4, "Hệ thống cập nhật thông tin phòng/sĩ số cho tất cả lớp trong file.")

add_heading(doc, "3.5. Quản lý Đăng ký Môn học (Giai đoạn 2)", 2)
add_body(doc, "Sau khi Giai đoạn 1 hoàn tất (đã có phòng và sĩ số cho tất cả lớp), quản trị viên mở Giai đoạn 2 tại trang /admin/enrollment.")

add_heading(doc, "3.5.1. Mở và đóng Giai đoạn 2", 3)
add_numbered_step(doc, 1, "Truy cập Admin → Quản lý đăng ký môn học hoặc URL /admin/enrollment.")
add_numbered_step(doc, 2, "Kiểm tra: tất cả lớp phải có phòng học và sĩ số tối đa (không có viền đỏ cảnh báo).")
add_numbered_step(doc, 3, "Nhấn Mở đăng ký HS để học sinh bắt đầu đăng ký môn học.")
add_numbered_step(doc, 4, "Sau thời hạn đăng ký, nhấn Đóng đăng ký HS.")
add_note_box(doc,
    "Khi Giai đoạn 2 đang mở, học sinh có thể đăng ký VÀ huỷ đăng ký. Sau khi đóng, không ai có thể thay đổi đăng ký.",
    title="Lưu ý quan trọng",
    color="FCE4D6",
    title_color=(0xC0, 0x50, 0x0)
)

add_heading(doc, "3.5.2. Xem danh sách học sinh theo lớp", 3)
add_body(doc, "Trang hiển thị các lớp đã publish dưới dạng bảng. Mỗi dòng hiển thị thông tin:")
add_table(doc,
    ["Thông tin", "Mô tả"],
    [
        ["#ID", "Mã số lớp trong hệ thống"],
        ["Giáo viên", "Họ tên giáo viên phụ trách"],
        ["Môn / Nhóm", "Tên môn học"],
        ["Khối", "Khối dành cho lớp (10/11/12)"],
        ["Thời gian", "Định dạng: Thứ X - Sáng/Chiều - Tiết X–Y"],
        ["Địa điểm", "Phòng học đã phân công"],
        ["Sĩ số", "Số HS đã đăng ký / Sĩ số tối đa (VD: 18/30). Cập nhật tự động mỗi 5 giây"],
        ["Thao tác", "Nút Xem HS – mở modal danh sách học sinh đã đăng ký"],
    ],
    col_widths=[1.5, 3.5, 3.0, 1.5, 4.0, 2.5, 1.5, 2.0]
)
add_body(doc, "Khi nhấn Xem HS, modal hiển thị bảng gồm: STT, Họ tên, CCCD, Khối, Lớp của từng học sinh đã đăng ký vào lớp đó.")
add_note_box(doc,
    "Chỉ những lớp đã được điền đầy đủ Địa điểm và Sĩ số tối đa mới hiển thị ở trang này "
    "và trong trang đăng ký của học sinh. Lớp thiếu thông tin sẽ không được học sinh nhìn thấy.",
    title="Điều kiện hiển thị lớp",
    color="DEEAF1",
    title_color=(0x1F, 0x49, 0x7D)
)

add_heading(doc, "3.5.3. Xuất kết quả đăng ký", 3)
add_numbered_step(doc, 1, "Sau khi đóng Giai đoạn 2, nhấn nút Xuất Excel kết quả.")
add_numbered_step(doc, 2, "File .xlsx được tải xuống, chứa danh sách từng lớp và học sinh đã đăng ký.")
add_numbered_step(doc, 3, "Dùng file này để lập danh sách lớp học, in sơ đồ chỗ ngồi, hoặc báo cáo.")
add_body(doc, "[Hình ảnh: Trang /admin/enrollment với accordion danh sách lớp và nút xuất Excel]", italic=True, color=(0x70, 0x70, 0x70))

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# PHẦN 4: GIÁO VIÊN
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "4. HƯỚNG DẪN SỬ DỤNG – GIÁO VIÊN", 1)
add_body(doc, "Giao diện dành cho giáo viên tập trung vào thời khoá biểu dạng grid (lưới), giúp giáo viên đăng ký mở lớp theo từng ô lịch trong tuần. URL sau đăng nhập: /teacher/schedule.")

add_heading(doc, "4.1. Tổng quan giao diện", 2)
add_body(doc, "Trang thời khoá biểu giáo viên gồm các thành phần chính:")
add_table(doc,
    ["Thành phần", "Vị trí", "Mô tả"],
    [
        ["Header", "Trên cùng", "Tên giáo viên, môn dạy, badge số lớp đã đăng ký (góc phải), nút Đăng xuất"],
        ["Bảng lịch tuần", "Giữa trang", "Grid: hàng = Buổi/Tiết, cột = Thứ 2 đến Thứ 7"],
        ["Ô lịch trống", "Trong bảng", "Màu xám nhạt – click để đăng ký mở lớp (khi Giai đoạn 1 mở)"],
        ["Ô lịch đã đăng ký", "Trong bảng", "Hiển thị: khối, số tiết, số HS đăng ký, phòng học"],
        ["Nút X huỷ lớp", "Góc ô lịch", "Hiện khi hover vào ô đã đăng ký (chỉ trong Giai đoạn 1)"],
        ["Badge số lớp", "Góc phải header", "Số lớp giáo viên đã đăng ký trong tuần này"],
    ],
    col_widths=[4.0, 3.0, 9.5]
)
add_body(doc, "[Hình ảnh: Giao diện thời khoá biểu giáo viên dạng grid với các ô lịch đã đăng ký]", italic=True, color=(0x70, 0x70, 0x70))

add_heading(doc, "4.2. Trạng thái giao diện theo giai đoạn", 2)
add_table(doc,
    ["Trạng thái hệ thống", "Giáo viên thấy gì / làm được gì"],
    [
        ["Giai đoạn 1 đang MỞ", "Ô lịch trống có thể click để đăng ký lớp. Ô đã có lớp hiện nút X để xoá."],
        ["Giai đoạn 1 đã ĐÓNG", "Chỉ xem thời khoá biểu — không thể thêm hoặc xoá lớp."],
        ["Giai đoạn 2 đang MỞ", "Xem thời khoá biểu; badge số HS/phòng cập nhật trực tiếp mỗi 5 giây."],
        ["Giai đoạn 2 đã ĐÓNG", "Chỉ xem thời khoá biểu và số HS cuối cùng đăng ký."],
    ],
    col_widths=[5.0, 11.5]
)

add_heading(doc, "4.4. Xem thời khoá biểu", 2)
add_body(doc, "Ngay sau khi đăng nhập, giáo viên thấy lịch tuần hiện tại. Cấu trúc bảng lịch:")
add_table(doc,
    ["Chiều", "Giá trị"],
    [
        ["Cột (Thứ)", "Thứ 2, Thứ 3, Thứ 4, Thứ 5, Thứ 6, Thứ 7"],
        ["Hàng – Buổi", "Buổi Sáng, Buổi Chiều"],
        ["Hàng – Tiết", "Tiết 1, Tiết 2, Tiết 3, Tiết 4 (trong mỗi buổi)"],
    ],
    col_widths=[4.0, 12.5]
)
add_body(doc, "Mỗi ô trong lịch đại diện cho một khung giờ cụ thể. Khi ô đã có lớp đăng ký, ô đó hiển thị:")
add_bullet(doc, "Khối học sinh (10 / 11 / 12)")
add_bullet(doc, "Số tiết của lớp")
add_bullet(doc, "Số học sinh đã đăng ký / Sĩ số tối đa (ví dụ: 15/30)")
add_bullet(doc, "Phòng học (nếu đã được phân công bởi Admin)")

add_heading(doc, "4.5. Đăng ký mở lớp", 2)
add_note_box(doc,
    "Chức năng đăng ký mở lớp CHỈ hoạt động khi Quản trị viên đã MỞ Giai đoạn 1. "
    "Nếu Giai đoạn 1 chưa mở hoặc đã đóng, các ô lịch trống sẽ không thể click.",
    title="Điều kiện tiên quyết",
    color="DEEAF1",
    title_color=(0x1F, 0x49, 0x7D)
)
add_numbered_step(doc, 1, "Xác định ô lịch muốn mở lớp (thứ mấy, buổi sáng/chiều, tiết mấy).")
add_numbered_step(doc, 2, "Click vào ô lịch trống tương ứng trong bảng.")
add_numbered_step(doc, 3, "Modal (hộp thoại) đăng ký mở lớp xuất hiện với hai trường cần điền:")
add_bullet(doc, "Khối: chọn 10, 11 hoặc 12 từ dropdown.", indent_level=2)
add_bullet(doc, "Số tiết: nhập số tiết học của lớp đó.", indent_level=2)
add_numbered_step(doc, 4, "Nhấn Xác nhận / Đăng ký trong modal.")
add_numbered_step(doc, 5, "Ô lịch chuyển màu và hiển thị thông tin lớp vừa đăng ký.")
add_numbered_step(doc, 6, "Badge số lớp ở góc phải header tăng lên 1.")
add_body(doc, "[Hình ảnh: Modal đăng ký mở lớp với dropdown chọn khối và ô nhập số tiết]", italic=True, color=(0x70, 0x70, 0x70))
add_note_box(doc,
    "Một ô lịch chỉ có thể đăng ký một lớp. Nếu ô đã có lớp, click vào ô đó sẽ hiện thông tin chi tiết, "
    "không mở form đăng ký mới.",
    title="Lưu ý",
)

add_heading(doc, "4.6. Xoá lớp đã đăng ký", 2)
add_note_box(doc,
    "Chỉ có thể xoá lớp khi Giai đoạn 1 đang mở. Sau khi Giai đoạn 1 đóng, lớp đã đăng ký không thể xoá.",
    title="Điều kiện",
    color="DEEAF1",
    title_color=(0x1F, 0x49, 0x7D)
)
add_numbered_step(doc, 1, "Di chuyển chuột (hover) vào ô lịch đã có lớp đăng ký.")
add_numbered_step(doc, 2, "Nút X màu đỏ xuất hiện ở góc trên bên phải của ô.")
add_numbered_step(doc, 3, "Click nút X.")
add_numbered_step(doc, 4, "Hộp thoại xác nhận xuất hiện: nhấn OK / Xác nhận để xoá lớp.")
add_numbered_step(doc, 5, "Ô lịch chuyển về trạng thái trống. Badge số lớp giảm đi 1.")
add_note_box(doc,
    "Xoá lớp sẽ xoá TOÀN BỘ đăng ký học sinh vào lớp đó (nếu Giai đoạn 2 đã từng mở). "
    "Tuy nhiên trong quy trình thông thường, Giai đoạn 1 kết thúc trước khi Giai đoạn 2 mở.",
    title="Cảnh báo",
    color="FCE4D6",
    title_color=(0xC0, 0x50, 0x0)
)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# PHẦN 5: HỌC SINH
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "5. HƯỚNG DẪN SỬ DỤNG – HỌC SINH", 1)
add_body(doc, "Giao diện học sinh gồm hai tab chính: Thời khoá biểu và Đăng ký môn học. URL sau đăng nhập: /student/schedule.")

add_heading(doc, "5.1. Đặt mật khẩu lần đầu / Đổi mật khẩu bắt buộc", 2)
add_body(doc, "Có hai trường hợp học sinh bị yêu cầu đổi mật khẩu bắt buộc:")
add_bullet(doc, "Lần đầu đăng nhập (tài khoản mới được tạo).")
add_bullet(doc, "Sau khi Quản trị viên đặt lại mật khẩu tạm thời.")
add_body(doc, "Trong cả hai trường hợp, sau khi xác thực thành công hệ thống tự chuyển đến trang Đổi mật khẩu bắt buộc:")
add_numbered_step(doc, 1, "Đọc hộp thông tin màu xanh hiển thị yêu cầu mật khẩu.")
add_numbered_step(doc, 2, "Nhập mật khẩu mới tại ô Mật khẩu mới.")
add_numbered_step(doc, 3, "Nhập lại mật khẩu tại ô Xác nhận mật khẩu (ô hiển thị ✓ Trùng khớp khi đúng).")
add_numbered_step(doc, 4, "Nhấn Xác nhận đổi mật khẩu.")
add_numbered_step(doc, 5, "Hệ thống chuyển về trang chính của học sinh.")
add_note_box(doc,
    "Yêu cầu mật khẩu: tối thiểu 8 ký tự; chỉ dùng chữ hoa (A–Z), chữ thường (a–z), chữ số (0–9) "
    "và ký tự đặc biệt !@#$%^&*()-_+=[]{}|<>,.?/~; "
    "KHÔNG dùng dấu nháy đơn ('), nháy kép (\"), chấm phẩy (;) hoặc gạch chéo (\\).",
    title="Yêu cầu mật khẩu",
    color="DEEAF1",
    title_color=(0x1F, 0x49, 0x7D)
)
add_note_box(doc,
    "Sau khi đổi mật khẩu, dùng mật khẩu mới này cho tất cả lần đăng nhập tiếp theo. "
    "Nếu quên mật khẩu, liên hệ Quản trị viên để được cấp mật khẩu tạm thời mới.",
    title="Quan trọng",
    color="FCE4D6",
    title_color=(0xC0, 0x50, 0x0)
)

add_heading(doc, "5.2. Trạng thái giao diện theo giai đoạn", 2)
add_table(doc,
    ["Trạng thái hệ thống", "Học sinh thấy gì / làm được gì"],
    [
        ["Giai đoạn 2 chưa mở", "Chỉ xem Tab Thời khoá biểu. Nút Đăng ký bị ẩn hoặc vô hiệu hoá."],
        ["Giai đoạn 2 đang MỞ", "Có thể đăng ký và huỷ đăng ký môn học. Badge sĩ số cập nhật tự động mỗi 5 giây."],
        ["Giai đoạn 2 đã ĐÓNG", "Chỉ xem thời khoá biểu cá nhân — không thể đăng ký, huỷ hoặc thực hiện bất kỳ thay đổi nào."],
    ],
    col_widths=[5.0, 11.5]
)
add_note_box(doc,
    "Khi Giai đoạn 2 đã đóng, toàn bộ kết quả đăng ký được giữ nguyên và không thay đổi. "
    "Học sinh chỉ có thể xem lịch học cá nhân của mình.",
    title="Sau khi kết thúc đăng ký",
    color="E2EFDA",
    title_color=(0x37, 0x86, 0x10)
)

add_heading(doc, "5.3. Tab Thời khoá biểu", 2)
add_body(doc, "Tab Thời khoá biểu hiển thị lịch tuần cá nhân của học sinh — chỉ đọc, hiển thị các môn đã đăng ký.")
add_table(doc,
    ["Trạng thái", "Hiển thị"],
    [
        ["Chưa đăng ký môn nào", "Lịch trống + thông báo nhắc nhở: \"Bạn chưa đăng ký môn học nào. Vào tab Đăng ký môn học để bắt đầu.\""],
        ["Đã có một số môn", "Ô lịch tô màu với tên môn, giáo viên, phòng học"],
        ["Ô lịch có lớp đăng ký", "Hiện: Tên môn, Giáo viên, Phòng học (nếu đã được phân công)"],
    ],
    col_widths=[4.5, 12.0]
)
add_body(doc, "[Hình ảnh: Tab thời khoá biểu học sinh – lịch trống với thông báo nhắc nhở]", italic=True, color=(0x70, 0x70, 0x70))
add_body(doc, "Cấu trúc bảng lịch học sinh giống bảng giáo viên: Thứ 2–7, Buổi Sáng/Chiều, Tiết 1–4.")

add_heading(doc, "5.4. Tab Đăng ký môn học", 2)
add_body(doc, "Đây là tab chính để học sinh xem và chọn môn học. Tab hiển thị hai khu vực:")

p = doc.add_paragraph()
run = p.add_run("Khu vực 1 – Môn đã đăng ký (phía trên):")
set_font(run, bold=True)
add_bullet(doc, "Liệt kê các môn học sinh đã chọn dưới dạng thẻ (card) gọn.")
add_bullet(doc, "Mỗi thẻ hiển thị: Tên môn, Giáo viên, Lịch, Phòng học.")
add_bullet(doc, "Có nút Huỷ đăng ký trên mỗi thẻ.")

p = doc.add_paragraph()
run = p.add_run("Khu vực 2 – Danh sách lớp có thể đăng ký (phía dưới):")
set_font(run, bold=True)
add_bullet(doc, "Hiển thị tất cả lớp học đang mở phù hợp với khối của học sinh.")
add_bullet(doc, "Mỗi thẻ lớp chứa thông tin đầy đủ:")
add_table(doc,
    ["Thông tin trên thẻ lớp", "Ý nghĩa"],
    [
        ["Tên môn học", "Môn học giáo viên phụ trách"],
        ["Giáo viên", "Họ tên giáo viên dạy lớp"],
        ["Khối", "Khối dành cho lớp này (10/11/12)"],
        ["Lịch học", "Thứ + Buổi + Tiết cụ thể"],
        ["Phòng học", "Phòng đã được phân công (hoặc \"Chưa phân phòng\")"],
        ["Badge số HS / Sĩ số", "Màu xanh: bình thường | Màu vàng: gần đầy (≥2/3 sĩ số) | Màu xám: đã đầy"],
        ["Cảnh báo gần đầy / đầy", "Chỉ hiện với HS chưa đăng ký lớp đó"],
        ["Nút Đăng ký", "Nhấn để đăng ký vào lớp (hiện khi chưa đăng ký lớp này)"],
        ["Nút Huỷ đăng ký", "Nhấn để huỷ (hiện khi đã đăng ký lớp này)"],
    ],
    col_widths=[5.5, 11.0]
)
add_body(doc, "[Hình ảnh: Tab đăng ký môn học với khu vực môn đã đăng ký và danh sách lớp có badge màu]", italic=True, color=(0x70, 0x70, 0x70))

add_heading(doc, "5.5. Đăng ký môn học", 2)
add_note_box(doc,
    "Chức năng đăng ký CHỈ hoạt động khi Quản trị viên đã MỞ Giai đoạn 2. "
    "Ngoài thời gian này, các nút Đăng ký sẽ bị vô hiệu hoá.",
    title="Điều kiện tiên quyết",
    color="DEEAF1",
    title_color=(0x1F, 0x49, 0x7D)
)
add_numbered_step(doc, 1, "Đăng nhập vào hệ thống với tài khoản học sinh.")
add_numbered_step(doc, 2, "Chuyển sang tab Đăng ký môn học.")
add_numbered_step(doc, 3, "Xem danh sách lớp trong Khu vực 2. Lưu ý màu badge số học sinh:")
add_bullet(doc, "Xanh lá: còn nhiều chỗ.", indent_level=2)
add_bullet(doc, "Vàng: gần đầy (≥ 2/3 sĩ số đã đăng ký).", indent_level=2)
add_bullet(doc, "Xám: đã đầy — không thể đăng ký thêm.", indent_level=2)
add_numbered_step(doc, 4, "Chọn lớp muốn học và nhấn nút Đăng ký trên thẻ đó.")
add_numbered_step(doc, 5, "Hệ thống kiểm tra trùng lịch (xem mục 5.6) và sĩ số:")
add_bullet(doc, "Nếu không trùng và còn chỗ: đăng ký thành công. Thông báo xanh xuất hiện.", indent_level=2)
add_bullet(doc, "Nếu trùng lịch: hiện modal thông báo trùng (xem mục 5.6).", indent_level=2)
add_bullet(doc, "Nếu lớp đã đầy: nút Đăng ký bị vô hiệu hoá (badge xám).", indent_level=2)
add_numbered_step(doc, 6, "Lớp vừa đăng ký xuất hiện trong Khu vực 1 (Môn đã đăng ký) và trong Tab Thời khoá biểu.")

add_heading(doc, "5.6. Huỷ đăng ký môn học", 2)
add_note_box(doc,
    "Chỉ có thể huỷ đăng ký khi Giai đoạn 2 đang mở. Sau khi đóng, không thể thay đổi.",
    title="Điều kiện",
    color="DEEAF1",
    title_color=(0x1F, 0x49, 0x7D)
)
add_body(doc, "Có hai cách để huỷ đăng ký môn học:")

p = doc.add_paragraph()
run = p.add_run("Cách 1 – Từ Khu vực 1 (Môn đã đăng ký):")
set_font(run, bold=True)
add_numbered_step(doc, 1, "Trong tab Đăng ký môn học, tìm môn muốn huỷ ở Khu vực 1 phía trên.")
add_numbered_step(doc, 2, "Nhấn nút Huỷ đăng ký trên thẻ môn đó.")
add_numbered_step(doc, 3, "Xác nhận huỷ trong hộp thoại.")

p = doc.add_paragraph()
run = p.add_run("Cách 2 – Từ danh sách lớp (Khu vực 2):")
set_font(run, bold=True)
add_numbered_step(doc, 1, "Tìm thẻ lớp đã đăng ký (có nút Huỷ đăng ký màu đỏ thay vì nút Đăng ký).")
add_numbered_step(doc, 2, "Nhấn nút Huỷ đăng ký.")
add_numbered_step(doc, 3, "Xác nhận huỷ. Thẻ lớp chuyển về trạng thái \"chưa đăng ký\".")

add_heading(doc, "5.7. Xử lý trùng lịch", 2)
add_body(doc, "Khi học sinh đăng ký một lớp có lịch trùng với lớp đã đăng ký trước đó, hệ thống tự động phát hiện và:")
add_numbered_step(doc, 1, "Dừng quá trình đăng ký.")
add_numbered_step(doc, 2, "Hiển thị modal thông báo trùng lịch với đầy đủ thông tin lớp bị trùng:")
add_table(doc,
    ["Thông tin trong modal", "Mô tả"],
    [
        ["Môn học bị trùng", "Tên môn đã đăng ký trước đó"],
        ["Giáo viên", "GV phụ trách lớp đã đăng ký"],
        ["Lịch học", "Thứ, buổi, tiết của lớp bị trùng"],
    ],
    col_widths=[5.5, 11.0]
)
add_numbered_step(doc, 3, "Nhấn Đóng để quay lại, sau đó chọn lớp khác phù hợp với lịch.")
add_note_box(doc,
    "Hệ thống kiểm tra trùng theo Thứ + Buổi + Tiết. Hai lớp cùng thứ cùng tiết dù khác buổi vẫn có thể tồn tại cạnh nhau nếu cấu hình tiết học không chồng chéo.",
    title="Cơ chế kiểm tra",
    color="E2EFDA",
    title_color=(0x37, 0x86, 0x10)
)
add_body(doc, "[Hình ảnh: Modal thông báo trùng lịch với thông tin lớp bị trùng]", italic=True, color=(0x70, 0x70, 0x70))

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# PHẦN 6: TÌNH HUỐNG THƯỜNG GẶP
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "6. CÁC TÌNH HUỐNG THƯỜNG GẶP VÀ CÁCH XỬ LÝ", 1)

problems = [
    (
        "Học sinh đăng nhập báo \"Sai mật khẩu\"",
        "Học sinh nhập sai mật khẩu hoặc MSSV.",
        [
            "Kiểm tra lại MSSV (không có khoảng trắng thừa).",
            "Kiểm tra mật khẩu (phân biệt hoa/thường).",
            "Nếu quên mật khẩu: liên hệ Quản trị viên để reset về mặc định (MSSV).",
        ]
    ),
    (
        "Giáo viên không thể click ô lịch để đăng ký",
        "Giai đoạn 1 chưa mở hoặc đã đóng.",
        [
            "Liên hệ Quản trị viên kiểm tra trạng thái Giai đoạn 1.",
            "Quản trị viên vào /admin và nhấn Mở đăng ký GV nếu chưa mở.",
        ]
    ),
    (
        "Học sinh không thấy nút Đăng ký trên thẻ lớp",
        "Giai đoạn 2 chưa mở, hoặc lớp đã đầy.",
        [
            "Kiểm tra badge màu sắc: nếu xám = đã đầy, không thể đăng ký.",
            "Nếu badge không xám: liên hệ Quản trị viên kiểm tra Giai đoạn 2.",
        ]
    ),
    (
        "Ô lớp học có viền đỏ trong trang /admin/class-reg",
        "Lớp chưa được điền phòng học hoặc sĩ số tối đa.",
        [
            "Quản trị viên click vào tên giáo viên để mở rộng danh sách lớp.",
            "Điền phòng học và sĩ số tối đa cho lớp có viền đỏ.",
            "Lưu tự động (Enter hoặc blur). Viền đỏ biến mất sau khi lưu đủ thông tin.",
        ]
    ),
    (
        "Không xuất được file Excel",
        "Trình duyệt chặn tải xuống tự động.",
        [
            "Kiểm tra thanh địa chỉ trình duyệt xem có thông báo chặn tải xuống không.",
            "Cho phép trang web này tải xuống file.",
            "Thử lại hoặc dùng trình duyệt khác.",
        ]
    ),
    (
        "Nhập Excel báo lỗi",
        "File Excel không đúng định dạng hoặc thiếu cột.",
        [
            "Tải file mẫu mới nhất từ hệ thống.",
            "Đảm bảo không xoá hoặc đổi tên cột trong file mẫu.",
            "Kiểm tra dữ liệu: không để ô bắt buộc trống, MSSV/Mã GV không trùng.",
            "Xem báo cáo lỗi chi tiết hiển thị sau khi upload để biết dòng nào lỗi.",
        ]
    ),
    (
        "Học sinh thấy lịch trống dù đã đăng ký",
        "Đang xem tab Thời khoá biểu nhưng chưa có môn đăng ký, hoặc cần tải lại trang.",
        [
            "Nhấn F5 hoặc Ctrl+R để tải lại trang.",
            "Chuyển sang tab Đăng ký môn học kiểm tra Khu vực 1 (Môn đã đăng ký).",
            "Nếu Khu vực 1 trống: học sinh chưa đăng ký thành công, cần thực hiện lại.",
        ]
    ),
    (
        "Trang web tải chậm hoặc báo lỗi kết nối",
        "Sự cố mạng hoặc máy chủ.",
        [
            "Kiểm tra kết nối mạng nội bộ trường.",
            "Thử tải lại trang (F5).",
            "Báo cáo cho bộ phận kỹ thuật nếu lỗi kéo dài.",
        ]
    ),
]

for title, cause, steps in problems:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(f"Vấn đề: {title}")
    set_font(run, bold=True, size=13, color=(0x1F, 0x49, 0x7D))

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.5)
    run2 = p2.add_run(f"Nguyên nhân thường gặp: {cause}")
    set_font(run2, italic=True, size=12, color=(0x60, 0x60, 0x60))

    p3 = doc.add_paragraph()
    run3 = p3.add_run("Cách xử lý:")
    set_font(run3, bold=True, size=12)
    for step in steps:
        add_bullet(doc, step)

add_page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# PHẦN 7: PHỤ LỤC
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "7. PHỤ LỤC", 1)

add_heading(doc, "7.1. Bảng phím tắt và mẹo nhanh", 2)
add_table(doc,
    ["Thao tác", "Phím tắt / Mẹo"],
    [
        ["Lưu thông tin phòng/sĩ số (Admin)", "Nhấn Enter hoặc click ra ngoài ô (auto-save)"],
        ["Tải lại trang", "F5 hoặc Ctrl+R"],
        ["Đăng xuất nhanh", "Nhấn nút Đăng xuất ở góc phải header"],
        ["Tìm giáo viên trong danh sách", "Dùng Ctrl+F trên trình duyệt để tìm kiếm văn bản"],
        ["Mở file Excel vừa tải về", "Kiểm tra thư mục Downloads hoặc thanh tải xuống của trình duyệt"],
    ],
    col_widths=[7.0, 9.5]
)

add_heading(doc, "7.2. Bảng màu sắc trạng thái", 2)
add_table(doc,
    ["Màu / Hiển thị", "Ý nghĩa", "Áp dụng cho"],
    [
        ["Xanh lá (badge số HS)", "Còn nhiều chỗ trống", "Thẻ lớp – Tab đăng ký HS"],
        ["Vàng (badge số HS)", "Gần đầy: ≥ 2/3 sĩ số đã đăng ký", "Thẻ lớp – Tab đăng ký HS"],
        ["Xám (badge số HS)", "Đã đầy: không nhận thêm HS", "Thẻ lớp – Tab đăng ký HS"],
        ["Viền đỏ (ô lớp)", "Lớp thiếu phòng hoặc sĩ số tối đa", "Trang /admin/class-reg"],
        ["Xanh nhạt (ô input)", "Đã lưu thành công (auto-save)", "Ô phòng học / sĩ số – Admin"],
        ["Ô lịch tô màu (GV/HS)", "Đã có lớp đăng ký ở ô này", "Bảng thời khoá biểu"],
    ],
    col_widths=[4.5, 6.5, 5.5]
)

add_heading(doc, "7.3. Sơ đồ quy trình tổng quát", 2)
add_body(doc, "[Hình ảnh: Sơ đồ flowchart quy trình từ Giai đoạn 1 đến Giai đoạn 2]", italic=True, color=(0x70, 0x70, 0x70))
add_body(doc, "Quy trình vận hành hệ thống mynhc theo các bước sau đây:")
steps_flow = [
    "Quản trị viên tạo tài khoản giáo viên và học sinh (hoặc nhập Excel).",
    "Quản trị viên MỞ Giai đoạn 1.",
    "Giáo viên đăng nhập và đăng ký các ô lịch mở lớp.",
    "Quản trị viên ĐÓNG Giai đoạn 1.",
    "Quản trị viên điền phòng học và sĩ số tối đa cho từng lớp (trang /admin/class-reg).",
    "Quản trị viên kiểm tra không còn lớp nào có viền đỏ cảnh báo.",
    "Quản trị viên MỞ Giai đoạn 2.",
    "Học sinh đăng nhập và đăng ký các môn học yêu thích.",
    "Quản trị viên ĐÓNG Giai đoạn 2.",
    "Quản trị viên xuất file Excel kết quả đăng ký từ trang /admin/enrollment.",
    "Sử dụng file Excel để lập danh sách lớp, in phòng học, báo cáo Ban Giám hiệu.",
]
for i, s in enumerate(steps_flow, 1):
    add_numbered_step(doc, i, s)

add_heading(doc, "7.4. Thông tin liên hệ hỗ trợ kỹ thuật", 2)
add_table(doc,
    ["Hình thức", "Thông tin"],
    [
        ["Đơn vị hỗ trợ", "Bộ phận CNTT – Trường THPT Nguyễn Hữu Cầu"],
        ["Email", "[email hỗ trợ kỹ thuật của trường]"],
        ["Điện thoại", "[số điện thoại bộ phận kỹ thuật]"],
        ["Giờ hỗ trợ", "Thứ 2 – Thứ 6: 7:00 – 17:00"],
        ["Phiên bản tài liệu", "1.0 – Tháng 8/2026"],
    ],
    col_widths=[5.0, 11.5]
)

# Footer note
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("─" * 50)
set_font(run, size=11, color=(0x80, 0x80, 0x80))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Tài liệu này thuộc bản quyền của Trường THPT Nguyễn Hữu Cầu.")
set_font(run, italic=True, size=11, color=(0x60, 0x60, 0x60))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Phiên bản 1.1 – Năm học 2026–2027 – Hệ thống mynhc (ClassReg)")
set_font(run, italic=True, size=11, color=(0x60, 0x60, 0x60))

# ─────────────────────────────────────────────────────────────────────────────
# ADD HEADER & FOOTER
# ─────────────────────────────────────────────────────────────────────────────
add_header_footer(doc)

# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────
output_path = "/Users/ddang/Documents/Subject_Survey/ClassReg/Hướng_dẫn_sử_dụng_mynhc.docx"
doc.save(output_path)
print(f"Đã lưu file tại: {output_path}")
